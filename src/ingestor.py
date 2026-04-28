"""
Pipeline d'ingestion multi-formats → ChromaDB
Gère : PDF, DOCX, XLSX, Images (OCR), JSON, Markdown
"""
import json
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from tqdm import tqdm

from src.config import (
    CHROMA_PATH, CHROMA_COLLECTION, EMBEDDING_MODEL,
    CHUNK_SIZE, CHUNK_OVERLAP, SUPPORTED_EXTENSIONS
)


# ─────────────────────────────────────────────
# Parseurs par format
# ─────────────────────────────────────────────

def parse_pdf(path: Path) -> List[Document]:
    """Extrait texte + métadonnées page par page avec pdfplumber."""
    import pdfplumber
    docs = []
    with pdfplumber.open(path) as pdf:
        meta_base = {
            "source": path.name,
            "source_path": str(path),
            "source_type": "pdf",
            "total_pages": len(pdf.pages),
        }
        # Métadonnées PDF si disponibles
        if pdf.metadata:
            if pdf.metadata.get("Title"):
                meta_base["doc_title"] = pdf.metadata["Title"]
            if pdf.metadata.get("Author"):
                meta_base["author"] = pdf.metadata["Author"]
            if pdf.metadata.get("CreationDate"):
                meta_base["creation_date"] = pdf.metadata["CreationDate"]

        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            # Extraction des tableaux comme texte structuré
            tables = page.extract_tables()
            for table in tables:
                rows = []
                for row in table:
                    if row:
                        rows.append(" | ".join(str(c) if c else "" for c in row))
                if rows:
                    text += "\n\n[TABLEAU]\n" + "\n".join(rows) + "\n[/TABLEAU]"

            if text.strip():
                docs.append(Document(
                    page_content=text.strip(),
                    metadata={**meta_base, "page": i, "chunk_context": f"Page {i}/{len(pdf.pages)}"}
                ))
    return docs


def parse_docx(path: Path) -> List[Document]:
    """Extrait texte Word par section/paragraphe avec hiérarchie."""
    from docx import Document as DocxDocument
    doc = DocxDocument(path)
    docs = []
    current_section = "Document"
    buffer = []
    section_num = 1

    meta_base = {
        "source": path.name,
        "source_path": str(path),
        "source_type": "docx",
    }
    # Propriétés du document
    core = doc.core_properties
    if core.author:
        meta_base["author"] = core.author
    if core.title:
        meta_base["doc_title"] = core.title
    if core.modified:
        meta_base["last_modified"] = str(core.modified)

    def flush(section, buf, num):
        text = "\n".join(buf).strip()
        if text:
            docs.append(Document(
                page_content=text,
                metadata={**meta_base, "section": section, "section_num": num,
                           "chunk_context": f"Section : {section}"}
            ))

    for para in doc.paragraphs:
        style = para.style.name.lower() if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if "heading" in style or style.startswith("titre"):
            flush(current_section, buffer, section_num)
            buffer = []
            current_section = text
            section_num += 1
            buffer.append(f"[TITRE] {text}")
        else:
            buffer.append(text)

    flush(current_section, buffer, section_num)

    # Tables Word
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            docs.append(Document(
                page_content="[TABLEAU]\n" + "\n".join(rows) + "\n[/TABLEAU]",
                metadata={**meta_base, "section": "Tableau", "chunk_context": "Tableau du document"}
            ))
    return docs


def parse_xlsx(path: Path) -> List[Document]:
    """Extrait chaque feuille Excel comme un document séparé."""
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    docs = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        headers = None
        for i, row in enumerate(ws.iter_rows(values_only=True), 1):
            clean = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in clean):
                if i == 1:
                    headers = clean
                    rows.append("COLONNES : " + " | ".join(clean))
                else:
                    if headers:
                        row_dict = {headers[j]: clean[j] for j in range(min(len(headers), len(clean)))}
                        rows.append(", ".join(f"{k}: {v}" for k, v in row_dict.items() if v))
                    else:
                        rows.append(" | ".join(c for c in clean if c))

        if rows:
            docs.append(Document(
                page_content="\n".join(rows),
                metadata={
                    "source": path.name,
                    "source_path": str(path),
                    "source_type": "xlsx",
                    "sheet_name": sheet_name,
                    "chunk_context": f"Feuille Excel : {sheet_name}",
                    "total_rows": ws.max_row,
                }
            ))
    return docs


def parse_image(path: Path) -> List[Document]:
    """OCR sur image (PNG, JPG, TIFF) avec pytesseract."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        print(f"  ⚠️  pytesseract non installé, image ignorée : {path.name}")
        return []

    # Chemin Tesseract sur Windows (à ajuster si vous avez installé ailleurs)
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    img = Image.open(path)
    # Tente français + anglais
    try:
        text = pytesseract.image_to_string(img, lang="fra+eng")
    except Exception:
        text = pytesseract.image_to_string(img)

    text = text.strip()
    if not text:
        return []

    return [Document(
        page_content=text,
        metadata={
            "source": path.name,
            "source_path": str(path),
            "source_type": "image_scan",
            "ocr_engine": "tesseract",
            "image_size": f"{img.width}x{img.height}",
            "chunk_context": f"Document scanné : {path.name}",
        }
    )]


def parse_json(path: Path) -> List[Document]:
    """Aplatit un JSON en blocs textuels sémantiques."""
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    docs = []
    meta_base = {
        "source": path.name,
        "source_path": str(path),
        "source_type": "json",
    }

    def flatten_obj(obj: Any, prefix: str = "", depth: int = 0) -> str:
        """Convertit récursivement un objet JSON en texte lisible."""
        if isinstance(obj, dict):
            parts = []
            for k, v in obj.items():
                label = f"{prefix}.{k}" if prefix else k
                parts.append(flatten_obj(v, label, depth + 1))
            return "\n".join(p for p in parts if p)
        elif isinstance(obj, list):
            return "\n".join(flatten_obj(i, prefix, depth + 1) for i in obj if flatten_obj(i, prefix, depth + 1))
        else:
            return f"{prefix}: {obj}" if str(obj).strip() else ""

    # Traitement intelligent selon structure
    if isinstance(data, dict):
        # Itère sur les clés racines comme sections
        for key, value in data.items():
            if key.startswith("_"):
                continue
            if isinstance(value, list):
                for i, item in enumerate(value):
                    text = flatten_obj(item, key)
                    if text.strip():
                        item_id = item.get("id", str(i)) if isinstance(item, dict) else str(i)
                        docs.append(Document(
                            page_content=text.strip(),
                            metadata={**meta_base, "json_key": key, "json_index": item_id,
                                       "chunk_context": f"JSON — {key}[{item_id}]"}
                        ))
            elif isinstance(value, dict):
                text = flatten_obj(value, key)
                if text.strip():
                    docs.append(Document(
                        page_content=text.strip(),
                        metadata={**meta_base, "json_key": key, "chunk_context": f"JSON — section {key}"}
                    ))
    else:
        text = flatten_obj(data)
        if text.strip():
            docs.append(Document(
                page_content=text.strip(),
                metadata={**meta_base, "chunk_context": "JSON — document complet"}
            ))
    return docs


def parse_markdown(path: Path) -> List[Document]:
    """Découpe un Markdown par sections H1/H2."""
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()

    docs = []
    meta_base = {
        "source": path.name,
        "source_path": str(path),
        "source_type": "markdown",
    }

    # Extraction du titre principal et des métadonnées
    title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    if title_match:
        meta_base["doc_title"] = title_match.group(1).strip()

    # Extraction de la source SharePoint/Drive si mentionnée
    source_match = re.search(r'\*\*(SharePoint|Google Drive|Drive)[^\*]+\*\*', content)
    if source_match:
        meta_base["platform"] = source_match.group(1)

    # Découpe par sections H2 (##)
    sections = re.split(r'\n(?=## )', content)
    for section in sections:
        section = section.strip()
        if not section:
            continue

        # Extraire le titre de la section
        title_match = re.match(r'##\s+(.+)\n', section)
        section_title = title_match.group(1).strip() if title_match else "Introduction"

        # Nettoie le Markdown brut
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', section)   # bold
        clean = re.sub(r'\*(.+?)\*', r'\1', clean)          # italic
        clean = re.sub(r'`(.+?)`', r'\1', clean)            # code inline
        clean = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean)  # links

        if clean.strip():
            docs.append(Document(
                page_content=clean.strip(),
                metadata={**meta_base, "section": section_title,
                           "chunk_context": f"Section : {section_title}"}
            ))
    return docs


def parse_txt(path: Path) -> List[Document]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    if not text:
        return []
    return [Document(
        page_content=text,
        metadata={"source": path.name, "source_path": str(path),
                  "source_type": "txt", "chunk_context": path.stem}
    )]


# ─────────────────────────────────────────────
# Dispatcher
# ─────────────────────────────────────────────

PARSERS = {
    ".pdf":  parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".xls":  parse_xlsx,
    ".png":  parse_image,
    ".jpg":  parse_image,
    ".jpeg": parse_image,
    ".tiff": parse_image,
    ".json": parse_json,
    ".md":   parse_markdown,
    ".txt":  parse_txt,
}


def parse_file(path: Path) -> List[Document]:
    """Parse un fichier selon son extension."""
    ext = path.suffix.lower()
    parser = PARSERS.get(ext)
    if not parser:
        print(f"  ⚠️  Format non supporté : {ext}")
        return []
    try:
        docs = parser(path)
        return docs
    except Exception as e:
        print(f"  ❌ Erreur lors du parsing de {path.name} : {e}")
        return []

    # Chemin Tesseract sur Windows (à ajuster si vous avez installé ailleurs)
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# ─────────────────────────────────────────────
# Chunking
# ─────────────────────────────────────────────

def chunk_documents(docs: List[Document]) -> List[Document]:
    """Découpe les documents en chunks avec overlap, préserve les métadonnées."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
        keep_separator=True,
    )
    chunks = []
    for doc in docs:
        splits = splitter.split_documents([doc])
        for i, chunk in enumerate(splits):
            # Ajoute un identifiant unique par chunk
            uid = hashlib.md5(
                f"{chunk.metadata.get('source', '')}_{i}_{chunk.page_content[:50]}".encode()
            ).hexdigest()[:12]
            chunk.metadata["chunk_id"] = uid
            chunk.metadata["chunk_index"] = i
            chunk.metadata["chunk_total"] = len(splits)
            chunks.append(chunk)
    return chunks


# ─────────────────────────────────────────────
# Ingestion principale
# ─────────────────────────────────────────────

def get_embeddings():
    """Charge le modèle d'embeddings local (téléchargé une seule fois)."""
    print(f"  📦 Chargement du modèle d'embeddings : {EMBEDDING_MODEL}")
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def ingest_documents(docs_dir: Path = None, reset: bool = False) -> Dict:
    """
    Pipeline complet : parse → chunk → embed → ChromaDB.
    
    Args:
        docs_dir: Dossier contenant les documents
        reset: Si True, vide la base avant ingestion
    
    Returns:
        Statistiques d'ingestion
    """
    docs_dir = docs_dir or Path("data/documents")
    stats = {"files_found": 0, "files_ok": 0, "files_error": 0,
             "chunks_total": 0, "skipped": 0}

    # Lister les fichiers supportés
    all_files = [
        f for f in docs_dir.rglob("*")
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
        and not f.name.startswith(".")
    ]
    stats["files_found"] = len(all_files)

    if not all_files:
        print(f"⚠️  Aucun fichier trouvé dans {docs_dir}")
        return stats

    print(f"\n📂 {len(all_files)} fichier(s) trouvé(s) dans {docs_dir}")

    # Embeddings
    embeddings = get_embeddings()

    # ChromaDB
    chroma_kwargs = {
        "collection_name": CHROMA_COLLECTION,
        "embedding_function": embeddings,
        "persist_directory": str(CHROMA_PATH),
    }

    if reset:
        import shutil
        if CHROMA_PATH.exists():
            shutil.rmtree(CHROMA_PATH)
            print("  🗑️  Base vectorielle réinitialisée")

    vectorstore = Chroma(**chroma_kwargs)

    # Récupérer les sources déjà indexées
    existing = set()
    try:
        existing_data = vectorstore.get(include=["metadatas"])
        existing = {m.get("source", "") for m in existing_data["metadatas"]}
    except Exception:
        pass

    # Parser + ingérer chaque fichier
    all_chunks = []
    for file_path in tqdm(all_files, desc="📄 Parsing"):
        if file_path.name in existing and not reset:
            print(f"  ⏭️  {file_path.name} — déjà indexé, ignoré")
            stats["skipped"] += 1
            continue

        raw_docs = parse_file(file_path)
        if not raw_docs:
            stats["files_error"] += 1
            continue

        chunks = chunk_documents(raw_docs)
        all_chunks.extend(chunks)
        stats["files_ok"] += 1
        print(f"  ✅ {file_path.name} → {len(raw_docs)} section(s) → {len(chunks)} chunk(s)")

    # Indexation par batchs
    if all_chunks:
        print(f"\n🔢 Indexation de {len(all_chunks)} chunks dans ChromaDB...")
        batch_size = 100
        for i in tqdm(range(0, len(all_chunks), batch_size), desc="💾 Indexation"):
            batch = all_chunks[i:i + batch_size]
            texts = [c.page_content for c in batch]
            metadatas = [c.metadata for c in batch]
            ids = [c.metadata["chunk_id"] for c in batch]
            vectorstore.add_texts(texts=texts, metadatas=metadatas, ids=ids)

        stats["chunks_total"] = len(all_chunks)
        print(f"\n✅ Ingestion terminée : {stats['chunks_total']} chunks indexés")
    else:
        print("\n⚠️  Aucun nouveau chunk à indexer")

    return stats
